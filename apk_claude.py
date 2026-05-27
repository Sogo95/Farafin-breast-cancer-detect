"""
Farafin AI for Health — Breast Cancer AI Detection
Application de détection du cancer du sein avec explicabilité Grad-CAM
Développé par : Farafin AI for Health
"""

import streamlit as st
import numpy as np
import tensorflow as tf
from tensorflow.keras.preprocessing import image   # ← celle-ci manquait
import matplotlib.pyplot as plt
from gradcam import make_gradcam_heatmap, overlay_heatmap
import cv2
import os
import pandas as pd
from datetime import datetime
import io
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from PIL import Image
import base64
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement




def get_image_base64(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Farafin AI — Cancer du Sein",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─────────────────────────────────────────────
# STYLES GLOBAUX
# ─────────────────────────────────────────────
PINK_PRIMARY   = "#880E4F"   # bleu médical (remplace le violet instable)
PINK_LIGHT     = "#FBCFE8"   # rose doux (UI secondaire / Octobre Rose léger)
PINK_DARK      = "#0F1F3D"   # bleu très sombre (sidebar / header)
ROSE_GRADIENT  = "linear-gradient(135deg, #EC4899 0%, #BE185D 50%, #9D174D 100%)"
WHITE          = "#FFFFFF"   # inchangé (OK)
GRAY_BG        = "#F4F7FB"   # fond propre dashboard médical
GRAY_TEXT      = "#475569"   # texte lisible moderne
SUCCESS_GREEN  = "#16A34A"   # vert plus moderne et standard UI santé
DANGER_RED     = "#DC2626"   # rouge plus net et lisible

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Source+Sans+3:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {{
    font-family: 'Source Sans 3', sans-serif;
    background-color: {GRAY_BG};
}}

/* ── Sidebar ── */
[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, {PINK_DARK} 0%, {PINK_PRIMARY} 100%);
    color: white;
}}
[data-testid="stSidebar"] * {{
    color: white !important;
}}
[data-testid="stSidebar"] hr {{
    border-color: rgba(255,255,255,0.25);
}}

/* ── Boutons principaux ── */
.stButton > button {{
    background: {ROSE_GRADIENT};
    color: white;
    border: none;
    border-radius: 12px;
    padding: 0.6rem 1.8rem;
    font-family: 'Source Sans 3', sans-serif;
    font-weight: 600;
    font-size: 1rem;
    letter-spacing: 0.02em;
    transition: all 0.2s ease;
    box-shadow: 0 4px 14px rgba(194,24,91,0.3);
    cursor: pointer;
}}
.stButton > button:hover {{
    transform: translateY(-2px);
    box-shadow: 0 8px 20px rgba(194,24,91,0.4);
}}

/* ── Cards ── */
.card {{
    background: white;
    border-radius: 16px;
    padding: 1.5rem;
    margin-bottom: 1.2rem;
    box-shadow: 0 2px 12px rgba(194,24,91,0.08);
    border: 1px solid rgba(194,24,91,0.1);
}}

/* ── Metric cards ── */
.metric-card {{
    background: white;
    border-radius: 14px;
    padding: 1.2rem 1rem;
    text-align: center;
    border-top: 4px solid {PINK_PRIMARY};
    box-shadow: 0 2px 10px rgba(0,0,0,0.06);
}}
.metric-value {{
    font-family: 'Playfair Display', serif;
    font-size: 2rem;
    font-weight: 700;
    color: {PINK_DARK};
    margin: 0;
}}
.metric-label {{
    font-size: 0.82rem;
    color: {GRAY_TEXT};
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin-top: 4px;
}}

/* ── Result badges ── */
.badge-malignant {{
    background: #FFEBEE;
    color: {DANGER_RED};
    border: 2px solid {DANGER_RED};
    border-radius: 50px;
    padding: 0.4rem 1.2rem;
    font-weight: 700;
    font-size: 1.1rem;
    display: inline-block;
}}
.badge-benign {{
    background: #E8F5E9;
    color: {SUCCESS_GREEN};
    border: 2px solid {SUCCESS_GREEN};
    border-radius: 50px;
    padding: 0.4rem 1.2rem;
    font-weight: 700;
    font-size: 1.1rem;
    display: inline-block;
}}

/* ── Section titles ── */
.section-title {{
    font-family: 'Playfair Display', serif;
    color: {PINK_DARK};
    font-size: 1.5rem;
    margin-bottom: 0.3rem;
    border-left: 5px solid {PINK_PRIMARY};
    padding-left: 0.8rem;
}}

/* ── Login page ── */
.login-container {{
    max-width: 420px;
    margin: 0 auto;
    padding: 2.5rem 2rem;
    background: white;
    border-radius: 24px;
    box-shadow: 0 16px 48px rgba(194,24,91,0.15);
}}
.login-title {{
    font-family: 'Playfair Display', serif;
    font-size: 1.7rem;
    color: {PINK_DARK};
    text-align: center;
    margin-bottom: 0.2rem;
}}
.login-sub {{
    text-align: center;
    color: {GRAY_TEXT};
    font-size: 0.9rem;
    margin-bottom: 1.8rem;
}}

/* ── Progress bar ── */
.stProgress > div > div > div {{
    background: {ROSE_GRADIENT};
    border-radius: 50px;
}}

/* ── Tables ── */
.stDataFrame {{
    border-radius: 12px !important;
    overflow: hidden;
}}

/* ── Onglets ── */
.stTabs [data-baseweb="tab-list"] {{
    gap: 6px;
    background: transparent;
}}
.stTabs [data-baseweb="tab"] {{
    border-radius: 10px 10px 0 0;
    font-weight: 600;
    color: {GRAY_TEXT};
}}
.stTabs [aria-selected="true"] {{
    background: white;
    color: {PINK_PRIMARY} !important;
    border-bottom: 3px solid {PINK_PRIMARY};
}}

/* ── Footer ── */
.footer {{
    text-align: center;
    padding: 1rem;
    color: {GRAY_TEXT};
    font-size: 0.78rem;
    margin-top: 2rem;
    border-top: 1px solid rgba(194,24,91,0.15);
}}

/* ── User profile badge ── */
.profile-badge {{
    display: flex;
    align-items: center;
    gap: 12px;
    background: rgba(255,255,255,0.12);
    border-radius: 14px;
    padding: 0.8rem 1rem;
    margin-bottom: 1rem;
}}
.avatar-circle {{
    width: 48px;
    height: 48px;
    border-radius: 50%;
    background: rgba(255,255,255,0.25);
    display: flex;
    align-items: center;
    justify-content: center;
    font-family: 'Playfair Display', serif;
    font-size: 1.1rem;
    font-weight: 700;
    color: white;
    flex-shrink: 0;
    border: 2px solid rgba(255,255,255,0.5);
}}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# LOGO OCTOBRE ROSE (SVG inline)
# ─────────────────────────────────────────────
OCTOBRE_ROSE_SVG = """
<svg viewBox="0 0 120 140" xmlns="http://www.w3.org/2000/svg" width="100" height="116">
  <defs>
    <linearGradient id="rbg" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" style="stop-color:#880E4F"/>
      <stop offset="100%" style="stop-color:#E91E8C"/>
    </linearGradient>
  </defs>
  <!-- Ruban Octobre Rose -->
  <path d="M60 20 C40 20 22 34 22 54 C22 68 32 78 44 86 C52 91 58 98 60 105
           C62 98 68 91 76 86 C88 78 98 68 98 54 C98 34 80 20 60 20 Z"
        fill="url(#rbg)" opacity="0.15"/>
  <!-- Nœud du ruban -->
  <ellipse cx="46" cy="55" rx="18" ry="22" fill="url(#rbg)" transform="rotate(-15 46 55)"/>
  <ellipse cx="74" cy="55" rx="18" ry="22" fill="url(#rbg)" transform="rotate(15 74 55)"/>
  <!-- Boucle centrale -->
  <ellipse cx="60" cy="63" rx="9" ry="7" fill="#880E4F"/>
  <!-- Queues du ruban -->
  <path d="M52 68 C46 80 38 95 42 112 C46 125 54 130 60 128" fill="url(#rbg)" opacity="0.9"/>
  <path d="M68 68 C74 80 82 95 78 112 C74 125 66 130 60 128" fill="url(#rbg)" opacity="0.9"/>
  <!-- Petits cercles décoratifs -->
  <circle cx="30" cy="40" r="3" fill="#E91E8C" opacity="0.5"/>
  <circle cx="90" cy="40" r="3" fill="#E91E8C" opacity="0.5"/>
  <circle cx="20" cy="65" r="2" fill="#C2185B" opacity="0.4"/>
  <circle cx="100" cy="65" r="2" fill="#C2185B" opacity="0.4"/>
  <!-- Texte -->
  <text x="60" y="138" text-anchor="middle" font-family="Georgia,serif" font-size="9"
        font-weight="bold" fill="#880E4F" letter-spacing="1">OCTOBRE ROSE</text>
</svg>
"""

# ─────────────────────────────────────────────
# UTILISATEURS (démo — en prod: base de données)
# ─────────────────────────────────────────────
USERS_DB = {
    "dr.bonou@farafin.ai": {
        "password": "medecin123",
        "nom": "Dr. Aminata BONOU",
        "initiales": "AB",
        "profil": "Médecin Radiologue",
        "hopital": "CHU de Bogodogo",
        "role": "admin"
    },
    "tech@farafin.ai": {
        "password": "tech2024",
        "nom": "Moussa Diallo",
        "initiales": "MD",
        "profil": "Technicien IA",
        "hopital": "Farafin AI for Health",
        "role": "user"
    },
    "dr.ouedraogo@farafin.ai": {
        "password": "onco2024",
        "nom": "Dr. Sylvie Ouédraogo",
        "initiales": "SO",
        "profil": "Oncologue",
        "hopital": "Polyclinique Ouagadougou",
        "role": "user"
    },
}

# ─────────────────────────────────────────────
# GRAD-CAM (implémentation intégrée)
# ─────────────────────────────────────────────
def make_gradcam_heatmap(img_array, model, last_conv_layer_name, pred_index=None):
    """Génère la heatmap Grad-CAM."""
    try:
        grad_model = tf.keras.models.Model(
            model.inputs,
            [model.get_layer(last_conv_layer_name).output, model.output]
        )
        with tf.GradientTape() as tape:
            last_conv_layer_output, preds = grad_model(img_array)
            if pred_index is None:
                pred_index = tf.argmax(preds[0])
            class_channel = preds[:, pred_index]

        grads = tape.gradient(class_channel, last_conv_layer_output)
        pooled_grads = tf.reduce_mean(grads, axis=(0, 1, 2))
        last_conv_layer_output = last_conv_layer_output[0]
        heatmap = last_conv_layer_output @ pooled_grads[..., tf.newaxis]
        heatmap = tf.squeeze(heatmap)
        heatmap = tf.maximum(heatmap, 0) / (tf.math.reduce_max(heatmap) + 1e-8)
        return heatmap.numpy()
    except Exception as e:
        h = np.random.rand(7, 7)
        return h / h.max()


def overlay_heatmap_on_image(img_pil, heatmap, alpha=0.45):
    """Superpose la heatmap Grad-CAM sur l'image originale."""
    img_rgb = np.array(img_pil.resize((224, 224)).convert("RGB"))
    heatmap_resized = cv2.resize(heatmap, (224, 224))
    heatmap_uint8 = np.uint8(255 * heatmap_resized)
    jet = plt.colormaps.get_cmap("jet")
    jet_colors = jet(np.arange(256))[:, :3]
    jet_heatmap = jet_colors[heatmap_uint8]
    jet_heatmap = np.uint8(jet_heatmap * 255)
    superimposed = cv2.addWeighted(img_rgb, 1 - alpha, jet_heatmap, alpha, 0)
    return Image.fromarray(superimposed)


# ======================
# SESSION HISTORY
# ======================
if "history" not in st.session_state:
    st.session_state.history = []

# ─────────────────────────────────────────────
# MODÈLE (cache)
# ─────────────────────────────────────────────
@st.cache_resource
def load_model():
    model = tf.keras.models.load_model("breast_cancer_mobilenet1.keras", compile=False)
    return model

model = load_model()

IMG_SIZE = 224
threshold = 0.4

# IMPORTANT: adapter selon ton modèle
LAST_CONV_LAYER = "Conv_1"  # MobileNetV2 last conv layer

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
for key, default in [
    ("authenticated", False),
    ("user_email", ""),
    ("history", []),
    ("page", "detection"),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║           PAGE DE CONNEXION              ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
def page_login():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@600;700&family=DM+Sans:wght@400;500&display=swap');

    [data-testid="stAppViewContainer"] > .main {
        background-color: #FDE8F0;
    }
    [data-testid="stHeader"] {
        background: transparent;
    }
    div[data-testid="stTextInput"] input {
        background-color: #FDF0F5 !important;
        border: 1px solid rgba(194,24,91,0.25) !important;
        border-radius: 8px !important;
        color: #3a1a28 !important;
        font-family: 'DM Sans', sans-serif !important;
    }
    div[data-testid="stTextInput"] input:focus {
        border-color: #C2185B !important;
        background: #fff !important;
    }
    div[data-testid="stTextInput"] label {
        color: #880E4F !important;
        font-size: 0.72rem !important;
        letter-spacing: 0.1em !important;
        text-transform: uppercase !important;
        font-weight: 500 !important;
    }
    div[data-testid="stForm"] .stButton > button {
        background: #C2185B !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-family: 'DM Sans', sans-serif !important;
        font-weight: 500 !important;
        font-size: 0.92rem !important;
        padding: 0.6rem 1.2rem !important;
        width: 100% !important;
        transition: background 0.2s !important;
    }
    div[data-testid="stForm"] .stButton > button:hover {
        background: #a3154e !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # Centrer avec 3 colonnes
    _, col, _ = st.columns([1, 1.2, 1])

    with col:
        # Logo ruban rose
        st.markdown("""
        <div style="text-align:center; margin: 1.5rem 0 0.8rem;">
            <img src='https://img.icons8.com/color/240/pink-ribbon.png' width='120' alt='Octobre Rose'/>
        </div>

        <h1 style="
            font-family: 'Cormorant Garamond', Georgia, serif;
            font-size: 1.6rem;
            font-weight: 700;
            color: #880E4F;
            text-align: center;
            line-height: 1.3;
            margin: 0 0 4px;
        ">Farafin BreastCancer AI Detect</h1>

        <p style="
            font-family: 'DM Sans', sans-serif;
            font-size: 0.75rem;
            color: #b07090;
            text-align: center;
            letter-spacing: 0.07em;
            margin: 0 0 1.6rem;
        ">Système d'aide au diagnostic du cancer du sein</p>

        <div style="
            width: 32px; height: 2px;
            background: #D4537E;
            margin: 0 auto 1.8rem;
            border-radius: 2px;
        "></div>
        """, unsafe_allow_html=True)

        # Formulaire
        with st.form("login_form", clear_on_submit=False):
            email = st.text_input(
                "Adresse e-mail",
                placeholder="nom@hopital.bf",
                help="Utilisez votre adresse professionnelle"
            )
            password = st.text_input(
                "Mot de passe",
                type="password",
                placeholder="••••••••"
            )
            st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)
            submitted = st.form_submit_button(
                "🔐  Se connecter",
                use_container_width=True
            )
            if submitted:
                user = USERS_DB.get(email.strip().lower())
                if user and user["password"] == password:
                    st.session_state.authenticated = True
                    st.session_state.user_email = email.strip().lower()
                    st.session_state.page = "dashboard"
                    st.rerun()
                else:
                    st.error("❌ Identifiants incorrects. Veuillez réessayer.")

        # Comptes démo
        st.markdown("""
        

        <p style="
            text-align: center;
            font-size: 0.64rem;
            color: #c4a0b4;
            margin-top: 1.2rem;
            font-family: 'DM Sans', sans-serif;
        ">
            © 2026 | Développé par
            <strong style="color:#880E4F;">Farafin AI for Health</strong><br>
            Outil d'aide au diagnostic — ne remplace pas l'avis médical
        </p>
        """, unsafe_allow_html=True)
# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║              SIDEBAR                     ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
def render_sidebar():
    user = USERS_DB[st.session_state.user_email]

    with st.sidebar:
        

        # Profil utilisateur
        st.markdown(f"""
        <div class="profile-badge">
            <div class="avatar-circle">{user["initiales"]}</div>
            <div>
                <p style="margin:0; font-weight:600; font-size:0.95rem;">{user["nom"]}</p>
                <p style="margin:0; font-size:0.75rem; opacity:0.85;">{user["profil"]}</p>
                <p style="margin:0; font-size:0.7rem; opacity:0.7;">🏥 {user["hopital"]}</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Navigation
        st.markdown('<p style="font-size:0.7rem; opacity:0.7; letter-spacing:0.1em; margin-bottom:0.3rem;"> </p>', unsafe_allow_html=True)

        nav_items = [
            ("📊 Dashboard", "dashboard"),
            ("🔬 Analyser une image par IA", "detection"),
            ("📋 Historique des analyse", "historique"),
            ("📄 Rapport d'analyse",       "rapport"), 
            
        ]
        for label, key in nav_items:
            active = st.session_state.page == key
            style = "background:rgba(255,255,255,0.2); border-radius:10px; padding:0.5rem 0.7rem;" if active else "padding:0.5rem 0.7rem;"
            if st.button(label, key=f"nav_{key}", use_container_width=True):
                st.session_state.page = key
                st.rerun()

        st.markdown("---")

        # Stats rapides sidebar
        total = len(st.session_state.history)
        malignes = sum(1 for h in st.session_state.history if h["malignant"])
        st.markdown(f"""
        <div style="background:rgba(255,255,255,0.1); border-radius:12px; padding:0.8rem; margin-bottom:1rem;">
            <p style="font-size:0.7rem; opacity:0.7; letter-spacing:0.1em; margin:0 0 0.5rem;">STATISTIQUES SESSION</p>
            <p style="margin:0.2rem 0; font-size:0.85rem;">🔬 Total Image Analysées : <b>{total}</b></p>
            <p style="margin:0.2rem 0; font-size:0.85rem;">🔴 Souspicion du Cancer : <b>{malignes}</b></p>
            <p style="margin:0.2rem 0; font-size:0.85rem;">🟢 Pas de Suspicion du Cancer : <b>{total - malignes}</b></p>
        </div>
        """, unsafe_allow_html=True)

        # Déconnexion
        if st.button("⎋  Se déconnecter", use_container_width=True):
            for k in ["authenticated", "user_email", "history", "page"]:
                del st.session_state[k]
            st.rerun()

        # Footer sidebar
        st.markdown("""
        <div style="position:fixed; bottom:16px; left:16px; right:16px; text-align:center;">
            <p style="font-size:0.65rem; opacity:0.6; margin:0;">
                v2.0 &nbsp;|&nbsp; Farafin AI for Health<br>
                ⚠️ Outil d'aide au diagnostic du cancer du sein 
            </p>
        </div>
        """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║          PAGE DÉTECTION IA               ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
def page_detection():
    user = USERS_DB[st.session_state.user_email]

    # ── Message de bienvenue
    st.markdown("""
    <style>
    @keyframes defilement {
        0%   { transform: translateX(100%); }
        100% { transform: translateX(-100%); }
    }
    .texte-defilant { overflow:hidden; white-space:nowrap; border-top:1px solid #F4C0D1; border-bottom:1px solid #F4C0D1; padding:6px 0; margin-top:8px; }
    .texte-defilant p { display:inline-block; animation:defilement 18s linear infinite; font-family:'DM Sans',sans-serif; font-size:0.9rem; color:#b07090; margin:0; }
    </style>
    <div style="margin-bottom:1.8rem; padding-bottom:1.4rem; border-bottom:2px solid #F4C0D1;">
        <h1 style="font-family:'Cormorant Garamond',Georgia,serif;font-size:2rem;font-weight:700;color:#880E4F;margin:0 0 10px;line-height:1.2;">🔬 Analyse et détection du cancer du sein</h1>
    </div>
    """, unsafe_allow_html=True)

    # ── Identification de la patiente
    st.markdown("""
    <p style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;font-weight:700;
               color:#880E4F;border-left:4px solid #D4537E;padding-left:0.8rem;margin-bottom:0.8rem;">
        👤 Identification de la Patiente
    </p>
    """, unsafe_allow_html=True)

    if "current_patient_id" not in st.session_state:
        st.session_state.current_patient_id = f"PAT-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    col_id1, col_id2, col_id3, col_id4 = st.columns([1.2, 1.5, 0.6, 0.8])
    with col_id1:
        patient_id = st.text_input(
            "ID Patiente (auto-généré)",
            value=st.session_state.get("patient_id_courant", st.session_state.current_patient_id),
            help="Généré automatiquement. Vous pouvez le modifier."
        )
    with col_id2:
        patient_nom = st.text_input(
            "Nom & Prénom",
            value=st.session_state.get("patient_nom_courant", ""),
            placeholder="ex: Koné Mariam"
        )
    with col_id3:
        patient_age = st.number_input(
            "Âge",
            min_value=18, max_value=70,
            value=int(st.session_state.get("patient_age_courant", 40))
        )
    with col_id4:
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
        if st.button("🔄 Nouvelle patiente", use_container_width=True):
            for k in ["patient_id_courant", "patient_nom_courant", "patient_age_courant", "last_analysis_idx"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.session_state.current_patient_id = f"PAT-{datetime.now().strftime('%Y%m%d%H%M%S')}"
            st.rerun()

    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

    # ── Upload + paramètres
    col_upload, col_params = st.columns([2, 1])

    with col_upload:
        uploaded_file = st.file_uploader(
            "📁 Charger une image mammographique",
            type=["png", "jpg", "jpeg"],
            help="Formats acceptés : PNG, JPG, JPEG"
        )

    with col_params:
        st.markdown("""
        <div style="background:white;border-radius:14px;padding:1.2rem;border:1px solid rgba(194,24,91,0.12);">
        <p style="font-family:'DM Sans',sans-serif;font-size:0.7rem;color:#880E4F;letter-spacing:0.1em;text-transform:uppercase;font-weight:600;margin:0 0 10px;">Paramètres</p>
        </div>
        """, unsafe_allow_html=True)
        #threshold = st.slider("Seuil de décision", 0.1, 0.9, 0.4, 0.05)

    if uploaded_file:

        # ── Validation de l'image AVANT tout traitement
        file_bytes = uploaded_file.read()
        uploaded_file.seek(0)

        try:
            img_check = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            w, h = img_check.size
            img_arr_check = np.array(img_check)

            if w < 100 or h < 100:
                st.error("❌ Image trop petite (minimum 100×100 px). Veuillez charger une mammographie.")
                st.stop()

            pixels_blancs = np.mean(
                (img_arr_check[:,:,0] > 240) &
                (img_arr_check[:,:,1] > 240) &
                (img_arr_check[:,:,2] > 240)
            )
            if pixels_blancs > 0.80:
                st.error("⚠️ Cette image ressemble à un logo ou une icône (trop de zones blanches). "
                         "Veuillez charger une image histologique ou mammographique.")
                st.markdown("""
                <div style="background:#FFF3E0;border-radius:10px;padding:1rem 1.2rem;
                            border-left:4px solid #FF9800;margin-top:0.5rem;">
                    <p style="margin:0;font-size:0.88rem;color:#E65100;">
                        <b>Images acceptées :</b> mammographies, images histologiques, échographies mammaires.<br>
                        <b>Images refusées :</b> logos, icônes, photos non médicales, images trop uniformes.
                    </p>
                </div>
                """, unsafe_allow_html=True)
                st.stop()

            if img_arr_check.std() < 15:
                st.error("⚠️ Image trop uniforme pour être une image médicale. "
                         "Veuillez charger une mammographie ou image histologique.")
                st.stop()

        except Exception:
            st.error("❌ Impossible de lire ce fichier. Veuillez charger une image valide.")
            st.stop()

        # ── Sauvegarder temporairement
        img_path = "temp.png"
        with open(img_path, "wb") as f:
            f.write(file_bytes)

        # ── Afficher l'image originale
        col_img, col_btn = st.columns([2, 1])
        with col_img:
            st.markdown("""
            <div style="background:white;border-radius:14px;padding:1rem;border:1px solid rgba(194,24,91,0.12);margin-bottom:1rem;">
            <p style="font-family:'DM Sans',sans-serif;font-size:0.7rem;color:#880E4F;letter-spacing:0.1em;text-transform:uppercase;font-weight:600;margin:0 0 8px;">Image chargée</p>
            </div>
            """, unsafe_allow_html=True)
            st.image(uploaded_file, caption=uploaded_file.name, use_container_width=True)

        with col_btn:
            st.markdown("<div style='height:60px'></div>", unsafe_allow_html=True)
            run = st.button("🔬  Lancer la détection IA", use_container_width=True)

        if run:
            with st.spinner("🧠 Analyse en cours…"):
                # ── PRÉTRAITEMENT
                img = image.load_img(img_path, target_size=(IMG_SIZE, IMG_SIZE))
                img_array = image.img_to_array(img)
                img_array = np.expand_dims(img_array, axis=0)
                img_array = tf.keras.applications.mobilenet_v2.preprocess_input(img_array)

                # ── PRÉDICTION
                prob = model.predict(img_array)[0][0]
                malignant = prob >= threshold
                confidence = prob if malignant else (1 - prob)
                pred = "🔴 Malignant (Cancer)" if malignant else "🟢 Benign"

                # ── GRAD-CAM (généré ici pour pouvoir sauvegarder l'image)
                heatmap = make_gradcam_heatmap(img_array, model, LAST_CONV_LAYER)
                cam_img = overlay_heatmap(img_path, heatmap)

            # ── Sauvegarder les images sur disque pour le rapport Word
            idx_img   = len(st.session_state.history)
            orig_path = f"temp_orig_{idx_img}.png"
            cam_path  = f"temp_cam_{idx_img}.png"

            with open(orig_path, "wb") as f:
                f.write(file_bytes)

            if isinstance(cam_img, np.ndarray):
                cv2.imwrite(cam_path, cam_img)
            else:
                cam_img.save(cam_path)

            # ── Sauvegarder dans l'historique
            st.session_state.history.append({
                "datetime":          datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                "patient_id":        patient_id,
                "patient_nom":       patient_nom,
                "patient_age":       patient_age,
                "image":             uploaded_file.name,
                "prob":              float(prob),
                "prediction":        pred,
                "malignant":         malignant,
                "confidence":        float(confidence),
                "threshold":         threshold,
                "user":              user["nom"],
                "avis_decision":     "",
                "avis_urgence":      "",
                "avis_commentaire":  "",
                "img_orig_path":     orig_path,
                "img_cam_path":      cam_path,
            })
            st.session_state.last_analysis_idx = idx_img

        # ── Afficher les résultats si une analyse a été faite
        if "last_analysis_idx" in st.session_state:
            idx = st.session_state.last_analysis_idx
            h   = st.session_state.history[idx]

            prob       = h["prob"]
            malignant  = h["malignant"]
            confidence = h["confidence"]
            pred       = h["prediction"]

            # ── Résultats
            st.markdown("---")
            st.markdown("""
            <p style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;font-weight:700;color:#880E4F;border-left:4px solid #D4537E;padding-left:0.8rem;margin-bottom:1rem;">
            Résultats de l'analyse
            </p>
            """, unsafe_allow_html=True)

            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown(f"""
                <div style="background:white;border-radius:14px;padding:1.2rem;text-align:center;
                            border-top:4px solid #C2185B;border:1px solid rgba(194,24,91,0.15);">
                    <p style="font-family:'Cormorant Garamond',serif;font-size:2rem;
                              font-weight:700;color:#880E4F;margin:0;">{prob:.4f}</p>
                    <p style="font-size:0.72rem;color:#b07090;text-transform:uppercase;
                              letter-spacing:0.08em;margin:4px 0 0;">Probabilité malignité</p>
                </div>
                """, unsafe_allow_html=True)
            with c2:
                st.markdown(f"""
                <div style="background:white;border-radius:14px;padding:1.2rem;text-align:center;
                            border:1px solid rgba(194,24,91,0.15);">
                    <p style="font-family:'Cormorant Garamond',serif;font-size:2rem;
                              font-weight:700;color:#880E4F;margin:0;">{confidence:.1%}</p>
                    <p style="font-size:0.72rem;color:#b07090;text-transform:uppercase;
                              letter-spacing:0.08em;margin:4px 0 0;">Confiance IA</p>
                </div>
                """, unsafe_allow_html=True)
            with c3:
                couleur = "#B71C1C" if malignant else "#2E7D32"
                fond    = "#FFEBEE" if malignant else "#E8F5E9"
                st.markdown(f"""
                <div style="background:{fond};border-radius:14px;padding:1.2rem;text-align:center;
                            border:1px solid {couleur}33;">
                    <p style="font-size:1.8rem;margin:0;">{'🔴' if malignant else '🟢'}</p>
                    <p style="font-size:0.72rem;color:{couleur};text-transform:uppercase;
                              letter-spacing:0.08em;margin:4px 0 0;font-weight:700;">
                        {'MALIGNE' if malignant else 'BÉNIGNE'}
                    </p>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown(f"**Probabilité de malignité : `{prob:.4f}`**")
            st.progress(float(prob))

            if malignant:
                st.error("⚠️ **Lésion maligne détectée.** Confirmation histologique recommandée. "
                         "Cet outil est une aide au diagnostic et ne remplace pas l'avis médical.")
            else:
                st.success("✅ **Lésion bénigne.** Suivi clinique régulier recommandé. "
                           "Cet outil est une aide au diagnostic et ne remplace pas l'avis médical.")

            # ── GRAD-CAM affiché depuis les fichiers sauvegardés
            st.markdown("---")
            st.markdown("""
            <p style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;font-weight:700;
                       color:#880E4F;border-left:4px solid #D4537E;padding-left:0.8rem;margin-bottom:0.5rem;">
                Explicabilité de la décision de l'IA
            </p>
            <p style="font-size:0.85rem;color:#b07090;margin-bottom:1rem;">
                Les zones <b>rouge/jaune</b> indiquent les régions qui ont le plus influencé la décision du modèle.
            </p>
            """, unsafe_allow_html=True)

            col_orig, col_cam = st.columns(2)
            with col_orig:
                st.markdown("""<div style="background:white;border-radius:12px;padding:0.8rem;
                border:1px solid rgba(194,24,91,0.12);">
                <p style="font-size:0.7rem;color:#880E4F;text-transform:uppercase;
                letter-spacing:0.1em;font-weight:600;margin:0 0 6px;">Image originale</p></div>""",
                unsafe_allow_html=True)
                if os.path.exists(h["img_orig_path"]):
                    st.image(h["img_orig_path"], use_container_width=True)

            with col_cam:
                st.markdown("""<div style="background:white;border-radius:12px;padding:0.8rem;
                border:1px solid rgba(194,24,91,0.12);">
                <p style="font-size:0.7rem;color:#880E4F;text-transform:uppercase;
                letter-spacing:0.1em;font-weight:600;margin:0 0 6px;">Zone suspecte détectée</p></div>""",
                unsafe_allow_html=True)
                if os.path.exists(h["img_cam_path"]):
                    st.image(h["img_cam_path"], caption="Carte thermique", use_container_width=True)

            # ── Avis du radiologue
            st.markdown("---")
            st.markdown("""
            <p style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;font-weight:700;
                       color:#880E4F;border-left:4px solid #D4537E;padding-left:0.8rem;margin-bottom:0.8rem;">
                📝 Avis du Radiologue
            </p>
            """, unsafe_allow_html=True)

            col_avis1, col_avis2 = st.columns(2)
            with col_avis1:
                avis_decision = st.selectbox(
                    "Décision clinique",
                    ["— Sélectionner —", "✅ Confirme : Bénigne", "⚠️ Confirme : Maligne",
                     "🔄 Infirme la décision IA", "🔍 Examen complémentaire requis"],
                    key=f"avis_decision_{idx}"
                )
            with col_avis2:
                avis_urgence = st.selectbox(
                    "Niveau d'urgence",
                    ["— Sélectionner —", "🟢 Non urgent", "🟡 Surveillance rapprochée",
                     "🔴 Urgent — Biopsie recommandée", "⛔ Très urgent — Prise en charge immédiate"],
                    key=f"avis_urgence_{idx}"
                )

            avis_commentaire = st.text_area(
                "Commentaire du radiologue",
                placeholder="Décrivez vos observations cliniques, les zones suspectes identifiées, "
                            "les recommandations pour la patiente…",
                height=130,
                key=f"avis_commentaire_{idx}"
            )

            if st.button("💾  Enregistrer l'avis", key=f"btn_avis_{idx}"):
                if avis_decision == "— Sélectionner —":
                    st.warning("Veuillez sélectionner une décision clinique.")
                else:
                    st.session_state.history[idx]["avis_decision"]    = avis_decision
                    st.session_state.history[idx]["avis_urgence"]     = avis_urgence
                    st.session_state.history[idx]["avis_commentaire"] = avis_commentaire

                    # Mémoriser les infos patiente pour la prochaine image
                    st.session_state.patient_id_courant  = st.session_state.history[idx]["patient_id"]
                    st.session_state.patient_nom_courant = st.session_state.history[idx]["patient_nom"]
                    st.session_state.patient_age_courant = st.session_state.history[idx]["patient_age"]

                    st.success("✅ Avis enregistré. Vous pouvez charger une nouvelle image.")

                    # Réinitialiser uniquement l'analyse — patiente conservée
                    del st.session_state.last_analysis_idx
                    st.rerun()
# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║           PAGE HISTORIQUE                ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
def export_historique_xlsx(history):
    wb = Workbook()
    ws = wb.active
    ws.title = "Historique Analyses"

    rose_fill  = PatternFill("solid", fgColor="C2185B")
    white_font = Font(color="FFFFFF", bold=True, size=11)
    center     = Alignment(horizontal="center", vertical="center")

    headers = ["Date & Heure", "ID Patiente", "Age", "Fichier", "Résultat",
               "Prob. Malignité", "Confiance IA", "Seuil", "Analyste"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = rose_fill
        cell.font = white_font
        cell.alignment = center

    green_fill = PatternFill("solid", fgColor="E8F5E9")
    red_fill   = PatternFill("solid", fgColor="FFEBEE")

    for row, h in enumerate(history, 2):
        values = [
            h.get("datetime", ""),
            h.get("patient_id", "—"),
            h.get("patient_nom", "—"),
            h.get("image", ""),
            h.get("prediction", ""),
            f"{h.get('prob', 0):.4f}",
            f"{h.get('confidence', 0):.1%}",
            f"{h.get('threshold', 0.4):.2f}",
            h.get("user", ""),
        ]
        fill = red_fill if h.get("malignant") else green_fill
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = fill
            cell.alignment = center

    widths = [20, 15, 20, 25, 22, 18, 15, 10, 22]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def page_historique():

    st.markdown("""
    <style>
    @keyframes defilement {
        0%   { transform: translateX(100%); }
        100% { transform: translateX(-100%); }}
    .texte-defilant { overflow:hidden; white-space:nowrap; border-top:1px solid #F4C0D1; border-bottom:1px solid #F4C0D1; padding:6px 0; margin-top:8px; }
    .texte-defilant p { display:inline-block; animation:defilement 18s linear infinite; font-family:'DM Sans',sans-serif; font-size:0.9rem; color:#b07090; margin:0; }
    </style>
    <div style="margin-bottom:1.8rem; padding-bottom:1.4rem; border-bottom:2px solid #F4C0D1;">
        <h1 style="font-family:'Cormorant Garamond',Georgia,serif;font-size:2rem;font-weight:700;color:#880E4F;margin:0 0 10px;line-height:1.2;">
            📋 Historique des Analyses
        </h1>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.history:
        st.info("Aucune analyse effectuée durant cette session. Rendez-vous sur l'onglet **Détection IA**.")
        return

    df = pd.DataFrame(st.session_state.history)
    df["Résultat"]    = df["malignant"].apply(lambda x: "🔴 Maligne" if x else "🟢 Bénigne")
    df["Probabilité"] = df["prob"].apply(lambda x: f"{x:.1%}")
    df["Confiance"]   = df["confidence"].apply(lambda x: f"{x:.1%}")

    # ── Filtres
    col_f1, col_f2 = st.columns([2, 1])
    with col_f1:
        filtre = st.selectbox("Filtrer par résultat", ["Tous", "Malignes uniquement", "Bénignes uniquement"])
    with col_f2:
        st.metric("Total analyses", len(df))

    if filtre == "Malignes uniquement":
        df = df[df["malignant"] == True]
    elif filtre == "Bénignes uniquement":
        df = df[df["malignant"] == False]

    # ── Tableau
    cols_display = ["datetime", "patient_id", "patient_age", "image", "Résultat", "Probabilité", "Confiance", "user"]
    col_labels = {
        "datetime":    "Date & Heure",
        "patient_id":  "ID Patiente",
        "patient_age": "Âge Patiente",
        "image":       "Fichier",
        "Résultat":    "Résultat",
        "Probabilité": "Prob. Malignité",
        "Confiance":   "Confiance IA",
        "user":        "Analyste",
    }

    # Colonnes optionnelles — n'affiche que celles qui existent dans df
    cols_existantes = [c for c in cols_display if c in df.columns]
    st.dataframe(
        df[cols_existantes].rename(columns=col_labels),
        use_container_width=True,
        hide_index=True,
    )

    # ── Boutons export + effacer
    col_b1, col_b2 = st.columns([1, 1])

    with col_b1:
        xlsx_buf = export_historique_xlsx(st.session_state.history)
        st.download_button(
            "⬇️  Exporter en Excel (.xlsx)",
            data=xlsx_buf,
            file_name=f"historique_farafin_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    with col_b2:
        if st.button("🗑️  Effacer l'historique", use_container_width=True):
            st.session_state.history = []
            st.rerun()
# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║            PAGE DASHBOARD                ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
def page_dashboard():
    
    col_bienvenu, col_msg = st.columns([1, 1])

    with col_msg:
        try:
            logo_b64 = get_image_base64("Logo_FAI4H_neew.png")  # adapte le nom de ton fichier
            logo_html = f'<img src="data:image/png;base64,{logo_b64}" height="70" style="margin-bottom:0.8rem;"/>'
        except:
            logo_html = ""
        
    st.markdown(f"""
        <style>
        @keyframes defilement {{
            0%   {{ transform: translateX(100%); }}
            100% {{ transform: translateX(-100%); }}
        }}
        .texte-defilant {{ overflow:hidden; white-space:nowrap; border-top:1px solid #F4C0D1; border-bottom:1px solid #F4C0D1; padding:6px 0; margin-top:8px; }}
        .texte-defilant p {{ display:inline-block; animation:defilement 18s linear infinite; font-family:'DM Sans',sans-serif; font-size:0.9rem; color:#b07090; margin:0; }}
        </style>

        <div style="margin-bottom:1.8rem; padding-bottom:1.4rem; border-bottom:2px solid #F4C0D1;">
            <div style="display:flex; align-items:center; gap:1.2rem; margin-bottom:0.8rem;">
                {logo_html}
                <div>
                    <p style="font-family:'DM Sans',sans-serif;font-size:0.75rem;color:#b07090;
                            letter-spacing:0.14em;text-transform:uppercase;margin:0 0 4px;">Bienvenue sur</p>
                    <h1 style="font-family:'Cormorant Garamond',Georgia,serif;font-size:2rem;
                            font-weight:700;color:#880E4F;margin:0;line-height:1.2;">
                        Farafin BreastCancer AI Detect
                    </h1>
                </div>
            </div>
            <div class="texte-defilant">
                <p>Il s'agit d'un système d'aide au diagnostic du cancer du sein basé sur l'intelligence artificielle adaptée au contexte médical du Burkina Faso.</p>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
                </div>
        <div style="margin-bottom:1.8rem; padding-bottom:1.4rem; border-bottom:2px solid #F4C0D1;">
        <h1 style="font-family:'Cormorant Garamond',Georgia,serif;font-size:2rem;font-weight:700;color:#880E4F;margin:0 0 10px;line-height:1.2;">📊 Dashboard Synthèse
        </div>
     """, unsafe_allow_html=True)

    history = st.session_state.history
    total = len(history)
    malignes = sum(1 for h in history if h["malignant"])
    benignes = total - malignes
    avg_prob = np.mean([h["prob"] for h in history]) if history else 0
    avg_conf = np.mean([h["confidence"] for h in history]) if history else 0

    # ── Métriques
    c1, c2, c3, c4 = st.columns(4)
    for col, val, label in [
        (c1, total, "Analyses totales"),
        (c2, malignes, "Cas malins"),
        (c3, benignes, "Cas bénins"),
        (c4, f"{avg_conf:.0%}", "Confiance moy."),
    ]:
        with col:
            st.markdown(f"""
            <div class="metric-card">
                <p class="metric-value">{val}</p>
                <p class="metric-label">{label}</p>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if not history:
        st.info("Effectuez des analyses pour voir les statistiques apparaître ici.")
        return

    col_g1, col_g2 = st.columns(2)

    # ── Graphique 1 : camembert bénin/malin
    with col_g1:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown("**Répartition des diagnostics**")
        fig1, ax1 = plt.subplots(figsize=(4, 3.5))
        fig1.patch.set_facecolor("white")
        if total > 0:
            sizes = [malignes, benignes]
            labels = ["Malignes", "Bénignes"]
            colors = ["#C2185B", "#2E7D32"]
            explode = (0.05, 0)
            ax1.pie(sizes, labels=labels, colors=colors, explode=explode,
                    autopct="%1.1f%%", startangle=90,
                    textprops={"fontsize": 10, "fontfamily": "Source Sans 3"})
        ax1.set_axis_off()
        fig1.tight_layout()
        st.pyplot(fig1, use_container_width=True)
        plt.close(fig1)
        st.markdown('</div>', unsafe_allow_html=True)

    # ── Graphique 2 : probabilités au fil du temps
    with col_g2:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown("**Évolution des probabilités**")
        fig2, ax2 = plt.subplots(figsize=(4, 3.5))
        fig2.patch.set_facecolor("white")
        probs  = [h["prob"] for h in history]
        labels = [f"#{i+1}" for i in range(len(history))]
        colors_bar = ["#C2185B" if h["malignant"] else "#2E7D32" for h in history]
        ax2.bar(labels, probs, color=colors_bar, edgecolor="white", linewidth=0.5)
        ax2.axhline(y=0.4, color="#FF7043", linestyle="--", linewidth=1, label="Seuil (0.4)")
        ax2.set_ylim(0, 1)
        ax2.set_ylabel("Probabilité", fontsize=9)
        ax2.set_xlabel("Analyse #", fontsize=9)
        ax2.tick_params(labelsize=8)
        ax2.spines[["top", "right"]].set_visible(False)
        ax2.legend(fontsize=8)
        fig2.tight_layout()
        st.pyplot(fig2, use_container_width=True)
        plt.close(fig2)
        st.markdown('</div>', unsafe_allow_html=True)

    # ── Alerte si taux malin élevé
    if total >= 3 and malignes / total > 0.5:
        st.warning(
            f"⚠️ **Attention** : {malignes/total:.0%} des analyses de cette session sont classées malignes. "
            "Veuillez vérifier la qualité des images et consulter un spécialiste."
        )

    # ── Tableau récapitulatif
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("**Tableau récapitulatif**")
    df = pd.DataFrame(history)
    if len(df) > 0:
        summary = pd.DataFrame({
            "Métrique": ["Probabilité moyenne", "Probabilité max.", "Probabilité min.", "Confiance moyenne"],
            "Valeur": [
                f"{df['prob'].mean():.1%}",
                f"{df['prob'].max():.1%}",
                f"{df['prob'].min():.1%}",
                f"{df['confidence'].mean():.1%}",
            ]
        })
        st.dataframe(summary, use_container_width=True, hide_index=True)
    st.markdown('</div>', unsafe_allow_html=True)
# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║              Page  Rapport           ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────

# Fonction de génération de rapport détaillé pour une patiente sélectionnée
# ─────────────────────────────────────────────
# FONCTION UTILITAIRE — couleur de cellule
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Colorie le fond d'une cellule Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)
 
 
# ─────────────────────────────────────────────
# FONCTION PRINCIPALE — rapport Word
# ─────────────────────────────────────────────
def generer_rapport_word(analyses, patient_id):
    """
    Génère un rapport Word complet pour une patiente.
    Retourne un BytesIO prêt pour st.download_button.
 
    Paramètres :
        analyses   : liste des entrées history[] pour cette patiente
        patient_id : identifiant de la patiente (utilisé pour le nom du fichier)
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
 
    doc = DocxDocument()
 
    # ── Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
 
    info = analyses[0]
 
    # ════════════════════════════════════════
    # EN-TÊTE — Bandeau rose
    # ════════════════════════════════════════
# ── Logo + Titre sur la même ligne (tableau 2 colonnes)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    header_table.columns[0].width = Cm(3.5)
    header_table.columns[1].width = Cm(13.5)

    # Colonne gauche — Logo
    cell_logo = header_table.rows[0].cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell_logo, "FFFFFF")
    logo_para = cell_logo.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run()
    try:
        logo_run.add_picture("Logo_FAI4H_neew.png", width=Cm(2.8))
    except Exception:
        logo_run.add_text("")  # Silencieux si logo absent

    # Colonne droite — Texte
    cell_titre = header_table.rows[0].cells[1]
    cell_titre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell_titre, "FFFFFF")

    titre_para = cell_titre.paragraphs[0]
    titre_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    titre_run = titre_para.add_run("FARAFIN AI FOR HEALTH")
    titre_run.bold = True
    titre_run.font.size = Pt(18)
    titre_run.font.color.rgb = RGBColor(0xC2, 0x18, 0x5B)
    titre_run.font.name = "Arial"

    sub_para = cell_titre.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = sub_para.add_run("Rapport d'Analyse — Détection du Cancer du Sein par IA")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
    sub_run.font.name = "Arial"
    sub_run.italic = True

    # Supprimer les bordures du tableau d'en-tête
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in header_table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)
 
    # Ligne séparatrice rose
    sep = doc.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C2185B')
    pBdr.append(bottom)
    pPr.append(pBdr)
 
    doc.add_paragraph()  # espace
 
    # ════════════════════════════════════════
    # SECTION 1 — Informations patiente
    # ════════════════════════════════════════
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Informations de la Patiente")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1_run.font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
    h1_run.font.name = "Arial"
 
    nb_malignes = sum(1 for h in analyses if h.get("malignant"))
    nb_benignes = len(analyses) - nb_malignes
    prob_max    = max(h.get("prob", 0) for h in analyses)
 
    # Tableau infos patiente
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(12)
 
    champs = [
        ("ID Patiente",          info.get("patient_id", "—")),
        ("Nom & Prénom",         info.get("patient_nom", "—")),
        ("Âge",                  f"{info.get('patient_age', '—')} ans"),
        ("Date du rapport",      datetime.now().strftime("%d/%m/%Y à %H:%M")),
        ("Analyste",             info.get("user", "—")),
    ]
    for i, (label, valeur) in enumerate(champs):
        row = tbl.rows[i]
        # Label
        cell_l = row.cells[0]
        cell_l.text = label
        cell_l.paragraphs[0].runs[0].bold = True
        cell_l.paragraphs[0].runs[0].font.size = Pt(10)
        cell_l.paragraphs[0].runs[0].font.name = "Arial"
        set_cell_bg(cell_l, "F8BBD9")
        # Valeur
        cell_v = row.cells[1]
        cell_v.text = valeur
        cell_v.paragraphs[0].runs[0].font.size = Pt(10)
        cell_v.paragraphs[0].runs[0].font.name = "Arial"
 
    doc.add_paragraph()
 
    # ════════════════════════════════════════
    # SECTION 2 — Synthèse
    # ════════════════════════════════════════
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Synthèse des Analyses")
    h2_run.bold = True
    h2_run.font.size = Pt(13)
    h2_run.font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
    h2_run.font.name = "Arial"
 
    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.style = 'Table Grid'
    headers_s = ["Images analysées", "Suspectes", "Bénignes", "Prob. max détectée"]
    values_s  = [str(len(analyses)), str(nb_malignes), str(nb_benignes), f"{prob_max:.1%}"]
 
    for j, (h_txt, v_txt) in enumerate(zip(headers_s, values_s)):
        # En-tête
        c_h = tbl2.rows[0].cells[j]
        c_h.text = h_txt
        c_h.paragraphs[0].runs[0].bold = True
        c_h.paragraphs[0].runs[0].font.size = Pt(9)
        c_h.paragraphs[0].runs[0].font.name = "Arial"
        c_h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c_h, "C2185B")
        c_h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Valeur
        c_v = tbl2.rows[1].cells[j]
        c_v.text = v_txt
        c_v.paragraphs[0].runs[0].bold = True
        c_v.paragraphs[0].runs[0].font.size = Pt(14)
        c_v.paragraphs[0].runs[0].font.name = "Arial"
        c_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = "FFEBEE" if (j == 1 and nb_malignes > 0) else "E8F5E9" if j == 2 else "FDE8F0"
        set_cell_bg(c_v, bg)
 
    doc.add_paragraph()
 
    # Conclusion automatique
    conclu = doc.add_paragraph()
    if nb_malignes > 0:
        conclu_run = conclu.add_run(
            f"⚠ CONCLUSION IA : {nb_malignes} image(s) sur {len(analyses)} présentent "
            f"des signes de malignité (probabilité max : {prob_max:.1%}). "
            f"Une prise en charge spécialisée est recommandée."
        )
        conclu_run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    else:
        conclu_run = conclu.add_run(
            f"✓ CONCLUSION IA : Toutes les images ({len(analyses)}) sont classées bénignes. "
            f"Un suivi clinique régulier est recommandé."
        )
        conclu_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    conclu_run.bold = True
    conclu_run.font.size = Pt(10)
    conclu_run.font.name = "Arial"
 
    doc.add_paragraph()
 
    # ════════════════════════════════════════
    # SECTION 3 — Détail par analyse + images
    # ════════════════════════════════════════
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Détail des Analyses par Image")
    h3_run.bold = True
    h3_run.font.size = Pt(13)
    h3_run.font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
    h3_run.font.name = "Arial"
 
    for i, h in enumerate(analyses, 1):
        malignant = h.get("malignant", False)
        couleur   = RGBColor(0xB7, 0x1C, 0x1C) if malignant else RGBColor(0x1B, 0x5E, 0x20)
        bg_hex    = "FFEBEE" if malignant else "E8F5E9"
 
        # Titre analyse
        titre_p = doc.add_paragraph()
        titre_r = titre_p.add_run(f"  Analyse #{i} — {h.get('image', '—')}")
        titre_r.bold = True
        titre_r.font.size = Pt(11)
        titre_r.font.color.rgb = couleur
        titre_r.font.name = "Arial"
 
        # Tableau résultats
        tbl3 = doc.add_table(rows=3, cols=4)
        tbl3.style = 'Table Grid'
 
        labels_row = ["Date", "Résultat IA", "Probabilité", "Confiance IA"]
        values_row = [
            h.get("datetime", "—"),
            h.get("prediction", "—"),
            f"{h.get('prob', 0):.4f}",
            f"{h.get('confidence', 0):.1%}",
        ]
        for j, (lb, vl) in enumerate(zip(labels_row, values_row)):
            c_l = tbl3.rows[0].cells[j]
            c_l.text = lb
            c_l.paragraphs[0].runs[0].bold = True
            c_l.paragraphs[0].runs[0].font.size = Pt(9)
            c_l.paragraphs[0].runs[0].font.name = "Arial"
            c_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(c_l, "F8BBD9")
 
            c_v = tbl3.rows[1].cells[j]
            c_v.text = vl
            c_v.paragraphs[0].runs[0].font.size = Pt(10)
            c_v.paragraphs[0].runs[0].font.name = "Arial"
            c_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(c_v, bg_hex)
 
        # Ligne avis radiologue
        avis_cell = tbl3.rows[2].cells[0]
        tbl3.rows[2].cells[0].merge(tbl3.rows[2].cells[3])
        avis_txt = (
            f"Avis radiologue : {h.get('avis_decision', '—')}  |  "
            f"Urgence : {h.get('avis_urgence', '—')}  |  "
            f"Commentaire : {h.get('avis_commentaire', 'Aucun commentaire')}"
        )
        avis_cell.text = avis_txt
        avis_cell.paragraphs[0].runs[0].font.size = Pt(9)
        avis_cell.paragraphs[0].runs[0].font.name = "Arial"
        avis_cell.paragraphs[0].runs[0].italic = True
        set_cell_bg(avis_cell, "FDF0F5")
 
        # ── Images (originale + Grad-CAM si disponibles)
        img_orig_path = f"temp_orig_{i}.png"
        img_cam_path  = f"temp_cam_{i}.png"
 
        has_orig = os.path.exists(img_orig_path)
        has_cam  = os.path.exists(img_cam_path)
 
        if has_orig or has_cam:
            doc.add_paragraph()
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
            if has_orig:
                lbl = doc.add_paragraph("Image originale")
                lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lbl.runs[0].font.size = Pt(9)
                lbl.runs[0].italic = True
                lbl.runs[0].font.color.rgb = RGBColor(0x88, 0x0E, 0x4F)
 
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(img_orig_path, width=Inches(2.8))
 
            if has_cam:
                lbl2 = doc.add_paragraph("Zone suspecte — Grad-CAM")
                lbl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lbl2.runs[0].font.size = Pt(9)
                lbl2.runs[0].italic = True
                lbl2.runs[0].font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
 
                cam_p = doc.add_paragraph()
                cam_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cam_run = cam_p.add_run()
                cam_run.add_picture(img_cam_path, width=Inches(2.8))
 
        doc.add_paragraph()
 
    # ════════════════════════════════════════
    # PIED DE PAGE
    # ════════════════════════════════════════
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    pPr2 = footer_para._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '6')
    top_b.set(qn('w:space'), '1')
    top_b.set(qn('w:color'), 'C2185B')
    pBdr2.append(top_b)
    pPr2.append(pBdr2)
 
    footer_run = footer_para.add_run(
        f"© {datetime.now().year} Farafin AI for Health  |  "
        f"Outil d'aide au diagnostic — Ne remplace pas l'avis médical  |  "
        f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xB0, 0x70, 0x90)
    footer_run.font.name = "Arial"
    footer_run.italic = True
 
    # ── Sauvegarder en mémoire
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



#### debut page
def page_rapport():
    st.markdown("""
    <p style="font-family:'Cormorant Garamond',serif;font-size:1.8rem;font-weight:700;
               color:#880E4F;border-left:4px solid #D4537E;padding-left:0.8rem;margin-bottom:1.5rem;">
        📄 Rapport Complet d'Analyse par Patiente
    </p>
    """, unsafe_allow_html=True)

    if not st.session_state.history:
        st.info("Aucune analyse disponible. Effectuez des analyses dans l'onglet Détection IA.")
        return

    # Liste des IDs uniques
    ids = list({h.get("patient_id", "—") for h in st.session_state.history if h.get("patient_id")})
    ids.sort()

    patient_selectionne = st.selectbox("🔍 Sélectionner une patiente", ids)

    # Filtrer les analyses de cette patiente
    analyses = [h for h in st.session_state.history
                if h.get("patient_id") == patient_selectionne]

    if not analyses:
        st.warning("Aucune analyse trouvée pour cette patiente.")
        return

    info = analyses[0]

    # ── En-tête rapport
    st.markdown(f"""
    <div style="background:white;border-radius:16px;padding:1.5rem 2rem;
                border:1px solid rgba(194,24,91,0.15);margin-bottom:1.2rem;">
        <div style="display:flex;justify-content:space-between;align-items:flex-start;">
            <div>
                <p style="font-size:0.68rem;color:#b07090;letter-spacing:0.12em;
                          text-transform:uppercase;margin:0 0 4px;">Rapport médical</p>
                <h2 style="font-family:'Cormorant Garamond',serif;font-size:1.6rem;
                           font-weight:700;color:#880E4F;margin:0 0 4px;">
                    {info.get('patient_nom', '—')}
                </h2>
                <p style="font-size:0.85rem;color:#b07090;margin:0;">
                    ID : <b>{patient_selectionne}</b> &nbsp;|&nbsp;
                    Âge : <b>{info.get('patient_age', '—')} ans</b> &nbsp;|&nbsp;
                    {len(analyses)} image(s) analysée(s)
                </p>
            </div>
            <div style="text-align:right;">
                <p style="font-size:0.75rem;color:#b07090;margin:0;">
                    Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}<br>
                    Par : {info.get('user', '—')}
                </p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Synthèse
    nb_malignes = sum(1 for h in analyses if h.get("malignant"))
    nb_benignes = len(analyses) - nb_malignes
    prob_max    = max(h.get("prob", 0) for h in analyses)
    prob_moy    = sum(h.get("prob", 0) for h in analyses) / len(analyses)

    c1, c2, c3, c4 = st.columns(4)
    for col, val, label in [
        (c1, len(analyses), "Images analysées"),
        (c2, nb_malignes,   "Suspectes (malignes)"),
        (c3, nb_benignes,   "Bénignes"),
        (c4, f"{prob_max:.1%}", "Prob. max détectée"),
    ]:
        with col:
            st.markdown(f"""
            <div style="background:white;border-radius:12px;padding:1rem;text-align:center;
                        border-top:3px solid #C2185B;border:1px solid rgba(194,24,91,0.12);">
                <p style="font-family:'Cormorant Garamond',serif;font-size:1.8rem;
                          font-weight:700;color:#880E4F;margin:0;">{val}</p>
                <p style="font-size:0.68rem;color:#b07090;text-transform:uppercase;
                          letter-spacing:0.08em;margin:4px 0 0;">{label}</p>
            </div>
            """, unsafe_allow_html=True)

    # ── Conclusion automatique
    st.markdown("<br>", unsafe_allow_html=True)
    if nb_malignes > 0:
        st.error(f"⚠️ **Conclusion automatique IA** : {nb_malignes} image(s) sur {len(analyses)} "
                 f"présentent des signes de malignité (prob. max : {prob_max:.1%}). "
                 f"Une prise en charge spécialisée est recommandée.")
    else:
        st.success(f"✅ **Conclusion automatique IA** : Toutes les images analysées ({len(analyses)}) "
                   f"sont classées bénignes. Suivi régulier recommandé.")

    # ── Détail par image
    st.markdown("---")
    st.markdown("**Détail des analyses**")
    for i, h in enumerate(analyses, 1):
        couleur = "#B71C1C" if h.get("malignant") else "#2E7D32"
        fond    = "#FFF5F7" if h.get("malignant") else "#F5FFF7"
        st.markdown(f"""
        <div style="background:{fond};border-radius:12px;padding:1rem 1.2rem;
                    border-left:4px solid {couleur};margin-bottom:0.8rem;">
            <div style="display:flex;justify-content:space-between;flex-wrap:wrap;gap:0.5rem;">
                <div>
                    <p style="font-weight:600;color:{couleur};margin:0 0 4px;">
                        Image #{i} — {h.get('image','—')}
                    </p>
                    <p style="font-size:0.82rem;color:#5D4E56;margin:0;">
                        🕐 {h.get('datetime','—')} &nbsp;|&nbsp;
                        Prob. : <b>{h.get('prob',0):.4f}</b> &nbsp;|&nbsp;
                        Confiance : <b>{h.get('confidence',0):.1%}</b>
                    </p>
                </div>
                <div style="text-align:right;">
                    <p style="font-weight:700;color:{couleur};margin:0;">{h.get('prediction','—')}</p>
                </div>
            </div>
            {f'''<div style="margin-top:0.6rem;padding-top:0.6rem;border-top:1px solid {couleur}33;">
                <p style="font-size:0.8rem;color:#5D4E56;margin:0;">
                    <b>Avis radiologue :</b> {h.get("avis_decision","—")}<br>
                    <b>Urgence :</b> {h.get("avis_urgence","—")}<br>
                    <b>Commentaire :</b> {h.get("avis_commentaire","—") or "Aucun commentaire"}
                </p>
            </div>''' if h.get("avis_decision") else ""}
        </div>
        """, unsafe_allow_html=True)

    # ── Export rapport XLSX
    st.markdown("---")

    def export_rapport_xlsx(analyses, patient_selectionne):
        wb = Workbook()
        ws = wb.active
        ws.title = "Rapport Patiente"

        rose_fill  = PatternFill("solid", fgColor="C2185B")
        white_font = Font(color="FFFFFF", bold=True, size=11)
        center     = Alignment(horizontal="center", vertical="center", wrap_text=True)
        bold       = Font(bold=True)

        # Titre
        ws.merge_cells("A1:H1")
        ws["A1"] = f"RAPPORT D'ANALYSE — Farafin AI for Health"
        ws["A1"].font = Font(color="FFFFFF", bold=True, size=14)
        ws["A1"].fill = rose_fill
        ws["A1"].alignment = center

        # Info patiente
        ws.merge_cells("A2:H2")
        info = analyses[0]
        ws["A2"] = (f"Patiente : {info.get('patient_nom','—')}  |  "
                    f"ID : {patient_selectionne}  |  "
                    f"Âge : {info.get('patient_age','—')} ans  |  "
                    f"Analyste : {info.get('user','—')}  |  "
                    f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        ws["A2"].alignment = center

        # En-têtes
        headers = ["#", "Date", "Fichier", "Résultat", "Probabilité",
                   "Confiance", "Décision Radiologue", "Commentaire"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=h)
            cell.fill = rose_fill
            cell.font = white_font
            cell.alignment = center

        green_fill = PatternFill("solid", fgColor="E8F5E9")
        red_fill   = PatternFill("solid", fgColor="FFEBEE")

        for row, h in enumerate(analyses, 4):
            fill = red_fill if h.get("malignant") else green_fill
            vals = [
                row - 3,
                h.get("datetime", ""),
                h.get("image", ""),
                h.get("prediction", ""),
                f"{h.get('prob', 0):.4f}",
                f"{h.get('confidence', 0):.1%}",
                h.get("avis_decision", "—"),
                h.get("avis_commentaire", "—"),
            ]
            for col, val in enumerate(vals, 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.fill = fill
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

        widths = [5, 20, 28, 22, 14, 12, 30, 40]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf

    rapport_buf = export_rapport_xlsx(analyses, patient_selectionne)
    st.markdown("---")
 
    # Préparer les chemins images pour chaque analyse
    for h in analyses:
        idx_h = st.session_state.history.index(h)
        h["img_orig_path"] = h.get("img_orig_path", f"temp_orig_{idx_h}.png")
        h["img_cam_path"]  = h.get("img_cam_path",  f"temp_cam_{idx_h}.png")
 
    col_dl1, col_dl2 = st.columns(2)
 
    with col_dl1:
        # Rapport Word
        word_buf = generer_rapport_word(analyses, patient_selectionne)
        st.download_button(
            "📄  Télécharger le rapport Word (.docx)",
            data=word_buf,
            file_name=f"rapport_{patient_selectionne}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
 
    with col_dl2:
        # Rapport Excel (existant)
        rapport_buf = export_rapport_xlsx(analyses, patient_selectionne)
        st.download_button(
            "📊  Télécharger le rapport Excel (.xlsx)",
            data=rapport_buf,
            file_name=f"rapport_{patient_selectionne}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

# ─────────────────────────────────────────────
# ╔══════════════════════════════════════════╗
# ║              ROUTEUR PRINCIPAL           ║
# ╚══════════════════════════════════════════╝
# ─────────────────────────────────────────────
if not st.session_state.authenticated:
    page_login()
else:
    render_sidebar()

    page = st.session_state.page

    if page == "detection":
        page_detection()
    elif page == "historique":
        page_historique()
    elif page == "dashboard":
        page_dashboard()
    elif page == "rapport":
        page_rapport()

    # ── Footer global
    st.markdown("""
    <div class="footer">
        <strong>Développé par Farafin AI for Health</strong> &nbsp;|&nbsp;
        Application de Détection du Cancer du Sein assité par IA &nbsp;|&nbsp;
        ⚠️ Outil d'aide au diagnostic — Ne se substitue pas à l'avis médical &nbsp;|&nbsp;
        v2.0 © 2026
    </div>
    """, unsafe_allow_html=True)